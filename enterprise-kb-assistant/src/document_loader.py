from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".md", ".markdown", ".txt", ".pdf", ".docx"}


def file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def load_file(path: Path, knowledge_base_id: str) -> list[Document]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持的文件格式：{suffix}")
    if suffix in {".md", ".markdown", ".txt"}:
        return _load_text(path, knowledge_base_id)
    if suffix == ".pdf":
        return _load_pdf(path, knowledge_base_id)
    if suffix == ".docx":
        return _load_docx(path, knowledge_base_id)
    raise ValueError(f"不支持的文件格式：{suffix}")


def _load_text(path: Path, knowledge_base_id: str) -> list[Document]:
    text = path.read_text(encoding="utf-8")
    return [
        Document(
            page_content=_clean(text),
            metadata={
                "knowledge_base_id": knowledge_base_id,
                "source": str(path),
                "file_name": path.name,
                "title": _title_from_text(text, path.stem),
                "page": None,
            },
        )
    ]


def _load_pdf(path: Path, knowledge_base_id: str) -> list[Document]:
    reader = PdfReader(str(path))
    docs: list[Document] = []
    for page_no, page in enumerate(reader.pages, start=1):
        text = _clean(page.extract_text() or "")
        if text:
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "knowledge_base_id": knowledge_base_id,
                        "source": str(path),
                        "file_name": path.name,
                        "title": path.stem,
                        "page": page_no,
                    },
                )
            )
    return docs


def _load_docx(path: Path, knowledge_base_id: str) -> list[Document]:
    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_text = _docx_tables_to_text(doc)
    text = "\n".join([*paragraphs, *table_text])
    return [
        Document(
            page_content=_clean(text),
            metadata={
                "knowledge_base_id": knowledge_base_id,
                "source": str(path),
                "file_name": path.name,
                "title": _title_from_text(text, path.stem),
                "page": None,
            },
        )
    ]


def _docx_tables_to_text(doc: DocxDocument) -> list[str]:
    tables: list[str] = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [_clean(cell.text).replace("\n", " / ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            tables.append(f"表格 {table_index}\n" + "\n".join(rows))
    return tables


def _clean(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r", "\n").split("\n")]
    return "\n".join(line for line in lines if line)


def _title_from_text(text: str, fallback: str) -> str:
    for line in text.splitlines():
        stripped = line.strip("# \t")
        if stripped:
            return stripped[:80]
    return fallback
